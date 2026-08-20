from fastapi import APIRouter, Depends, HTTPException, status, Query, Form, UploadFile, File
import os
from sqlalchemy.orm import Session
from database import get_db
from models.ticket import Ticket, SubTask
from models.user import User
from models.worklog import WorkLog
from models.project import Project
from models.team import Team
from models.audit_log import AuditLog
from schemas.ticket import TicketCreate, TicketUpdate, TicketResponse
from core.security import get_current_user
from typing import List, Optional
from datetime import date, timedelta, datetime
import shutil
from google import genai 
import traceback
from models.notification import Notification
from models.time_log import TimeLog
from models.daily_report import DailyReport
from models.comment import Comment
from models.task_type import TaskType  
from docx import Document
from fastapi.responses import StreamingResponse
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

router = APIRouter(
    prefix="/tickets",
    tags=["Tickets"]
)

def format_to_hhmm(hours_float: float) -> str:
    if not hours_float or hours_float <= 0:
        return "00:00"
    total_minutes = round(hours_float * 60)
    h = int(total_minutes // 60)
    m = int(total_minutes % 60)
    return f"{h:02d}:{m:02d}"

def log_action(db: Session, user_id: int, action: str, details: str, ticket_id: int = None):
    try:
        audit = AuditLog(
            user_id=user_id,
            action=action,
            details=details,
            ticket_id=ticket_id,
            created_at=datetime.now()
        )
        db.add(audit)
        db.commit()
    except Exception as e:
        print(f"Erro ao gravar log de auditoria: {e}")

def filter_tickets_by_permissions(query, current_user: User, db: Session):
    role = getattr(current_user, "role", "Member").lower()
    if role in ["admin", "manager", "gestor de operações"]:
        return query
    return query.outerjoin(SubTask, Ticket.id == SubTask.ticket_id).filter(
        (Ticket.assigned_to_id == current_user.id) | 
        (Ticket.creator_id == current_user.id) |
        (SubTask.assigned_to_id == current_user.id)
    ).distinct()

@router.get("/me/stats")
def get_my_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket)
    query = filter_tickets_by_permissions(query, current_user, db)
    
    all_tickets = query.all()
    today = date.today()
    
    to_do = sum(1 for t in all_tickets if t.status and t.status.lower() in ['to do', 'a fazer'])
    in_progress = sum(1 for t in all_tickets if t.status and t.status.lower() in ['in progress', 'em progresso'])
    in_review = sum(1 for t in all_tickets if t.status and t.status.lower() in ['in review', 'em revisão', 'em revisao'])
    done = sum(1 for t in all_tickets if t.status and t.status.lower() in ['done', 'concluído', 'concluido'])
    
    overdue = sum(1 for t in all_tickets if t.due_date and t.due_date < today and t.status and t.status.lower() not in ['done', 'concluído', 'concluido'])
    due_today = sum(1 for t in all_tickets if t.due_date and t.due_date == today and t.status and t.status.lower() not in ['done', 'concluído', 'concluido'])

    today_logs = db.query(TimeLog).filter(
        TimeLog.user_id == current_user.id,
        TimeLog.date == today
    ).all()
    hours_today = sum(log.hours_spent for log in today_logs)

    role = getattr(current_user, "role", "Member")

    return {
        "user_id": current_user.id,
        "role": role,
        "total_tickets": len(all_tickets),
        "to_do": to_do,
        "in_progress": in_progress,
        "in_review": in_review,
        "done": done,
        "overdue": overdue,
        "due_today": due_today,
        "hours_today": format_to_hhmm(hours_today)
    }

@router.get("/active", response_model=List[TicketResponse])
def get_active_tickets(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket).filter(Ticket.is_running == True)
    query = filter_tickets_by_permissions(query, current_user, db)
    return query.all()

@router.get("/statistics/chart-hours")
def get_chart_hours(
    period: str = Query("7"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    role = getattr(current_user, "role", "Member").lower()
    today = date.today()
    
    log_query = db.query(TimeLog)
    if role not in ["admin", "manager", "gestor de operações"]:
        log_query = log_query.filter(TimeLog.user_id == current_user.id)
        
    if str(period).lower() in ["7", "30", "1", "week", "last_week"]:
        num_days = 7 if str(period).lower() in ["7", "1", "week", "last_week"] else int(period)
        start_date = today - timedelta(days=num_days - 1)
        
        logs = log_query.filter(TimeLog.date >= start_date).all()
        
        days_array = [start_date + timedelta(days=i) for i in range(num_days)]
        hours_map = {d.isoformat(): 0.0 for d in days_array}
        
        for log in logs:
            if log.date:
                log_d = log.date.date() if hasattr(log.date, 'date') else log.date
                d_str = log_d.isoformat()
                if d_str in hours_map and log_d <= today:
                    hours_map[d_str] += (log.hours_spent or 0.0)
                    
        labels = [f"{d.strftime('%d/%m')}" for d in days_array]
        hours_data = [round(hours_map[d.isoformat()], 2) for d in days_array]
        
        return {"labels": labels, "hours": hours_data}
    else:
        start_month_date = today - timedelta(days=180)
        logs = log_query.filter(TimeLog.date >= start_month_date).all()
        
        months_keys = []
        for i in range(6):
            m_offset = 5 - i
            y = today.year
            m = today.month - m_offset
            while m <= 0:
                m += 12
                y -= 1
            months_keys.append((y, m))
            
        months_map = {f"{y}-{str(m).zfill(2)}": 0.0 for y, m in months_keys}
        
        for log in logs:
            if log.date:
                log_date_obj = log.date.date() if hasattr(log.date, 'date') else log.date
                key = f"{log_date_obj.year}-{str(log_date_obj.month).zfill(2)}"
                if key in months_map and log_date_obj <= today:
                    months_map[key] += (log.hours_spent or 0.0)
                    
        labels = [f"{str(m).zfill(2)}/{str(y)[2:]}" for y, m in months_keys]
        hours_data = [round(months_map[f"{y}-{str(m).zfill(2)}"], 2) for y, m in months_keys]
        
        return {"labels": labels, "hours": hours_data}

@router.get("/project/{project_id}", response_model=List[TicketResponse])
def get_project_tickets(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Projeto não encontrado.")
    
    tickets = db.query(Ticket).filter(Ticket.project_id == project_id).all()
    return tickets

@router.post("/", response_model=TicketResponse, status_code=status.HTTP_201_CREATED)
def create_ticket(
    ticket: TicketCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not ticket.description or ticket.description.strip() == "":
        raise HTTPException(status_code=400, detail="A descrição da tarefa é um campo obrigatório.")

    role = getattr(current_user, "role", "Member").lower()
    
    if role in ["user", "member", "técnico"]:
        assigned_id = current_user.id
        proj_id = None
    else:
        assigned_id = ticket.assigned_to_id
        proj_id = ticket.project_id if ticket.project_id else None

    blocked_id = getattr(ticket, 'blocked_by_id', None)
    if blocked_id:
        prereq = db.query(Ticket).filter(Ticket.id == blocked_id).first()
        if not prereq:
            raise HTTPException(status_code=400, detail="A tarefa antecedente especificada não existe.")

    s_date = ticket.start_date if ticket.start_date else None
    d_date = ticket.due_date if ticket.due_date else None

    db_ticket = Ticket(
        title=ticket.title,
        description=ticket.description,
        priority=ticket.priority,
        status=ticket.status,
        project_id=proj_id,
        client_id=ticket.client_id,
        assigned_to_id=assigned_id,
        estimated_hours=ticket.estimated_hours,
        due_date=d_date,
        start_date=s_date,
        blocked_by_id=blocked_id,
        creator_id=current_user.id 
    )
    db.add(db_ticket)
    db.commit()
    db.refresh(db_ticket)
    
    log_action(db, current_user.id, "Criação de Tarefa", f"Criou a tarefa #{db_ticket.id} - {db_ticket.title}", ticket_id=db_ticket.id)
    
    if db_ticket.assigned_to_id and db_ticket.assigned_to_id != current_user.id:
        notif = Notification(
            user_id=db_ticket.assigned_to_id,
            message=f"Foste atribuído a uma nova tarefa: {db_ticket.title}"
        )
        db.add(notif)
        db.commit()
        
    return db_ticket

@router.get("/", response_model=List[TicketResponse])
def get_tickets(
    search: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    check_and_create_deadline_notifications(current_user, db)
    query = db.query(Ticket)
    role = getattr(current_user, "role", "Member").lower()
    
    if role in ["user", "member", "técnico"]:
        query = query.outerjoin(SubTask, Ticket.id == SubTask.ticket_id).filter(
            (Ticket.assigned_to_id == current_user.id) |
            (Ticket.creator_id == current_user.id) |
            (SubTask.assigned_to_id == current_user.id)
        ).distinct()
    
    if search:
        query = query.filter(Ticket.title.ilike(f"%{search}%"))
    if status: 
        query = query.filter(Ticket.status == status)
        
    return query.all()

@router.get("/my-day/today")
def get_or_create_daily_report(
    target_date: Optional[str] = Query(None), 
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    if target_date:
        today = datetime.strptime(target_date, "%Y-%m-%d").date()
    else:
        today = date.today()
    
    report = db.query(DailyReport).filter(
        DailyReport.user_id == current_user.id,
        DailyReport.date == today
    ).first()
    
    if not report:
        report = DailyReport(
            user_id=current_user.id, 
            date=today, 
            status="Rascunho"
        )
        db.add(report)
        db.commit()
        db.refresh(report)
        
    logs = db.query(TimeLog).filter(
        TimeLog.user_id == current_user.id,
        TimeLog.date == today
    ).all()
    
    ticket_hours_map = {}
    for log in logs:
        if log.ticket_id and log.hours_spent > 0.001:
            ticket_hours_map[log.ticket_id] = ticket_hours_map.get(log.ticket_id, 0.0) + log.hours_spent
            
    tickets_data = []
    for ticket_id, total_hours in ticket_hours_map.items():
        t = db.query(Ticket).filter(Ticket.id == ticket_id).first()
        if t:
            tickets_data.append({
                "id": t.id,
                "title": t.title,
                "status": t.status,
                "priority": t.priority,
                "description": t.description,
                "hours_today": format_to_hhmm(total_hours)
            })
        
    return {
        "report": {
            "id": report.id,
            "status": report.status,
            "summary": report.summary,
            "detailed_report": report.detailed_report,
            "kilometers": report.kilometers,
            "overtime_hours": report.overtime_hours,
            "rejection_reason": getattr(report, "rejection_reason", None)
        },
        "tickets_worked": tickets_data
    }

@router.get("/my-day/week")
def get_week_status(
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    today = date.today()
    start_date = today - timedelta(days=6)
    
    reports = db.query(DailyReport).filter(
        DailyReport.user_id == current_user.id,
        DailyReport.date >= start_date,
        DailyReport.date <= today
    ).all()
    
    report_dict = {r.date: r for r in reports}
    week_data = []
    dias_semana = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]
    
    for i in range(7):
        current_d = start_date + timedelta(days=i)
        rep = report_dict.get(current_d)
        nome_dia = dias_semana[current_d.weekday()]
        
        if rep:
            status_rep = rep.status
        else:
            status_rep = "Em falta" if current_d < today else "Não iniciado"
            
        week_data.append({
            "date": current_d.isoformat(),
            "day_name": nome_dia,
            "day_num": current_d.day,
            "status": status_rep
        })
        
    return week_data

@router.put("/{ticket_id}", response_model=TicketResponse)
def update_ticket(
    ticket_id: int, 
    ticket_data: TicketUpdate, 
    current_user: User = Depends(get_current_user), 
    db: Session = Depends(get_db)
):
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket não encontrado")

    role = getattr(current_user, "role", "Member").lower()
    if role not in ["admin", "manager", "gestor de operações"]:
        creator_id = getattr(ticket, "creator_id", None)
        assigned_id = getattr(ticket, "assigned_to_id", None)
        if creator_id != current_user.id and assigned_id != current_user.id:
            raise HTTPException(status_code=403, detail="Acesso negado: Não podes editar esta tarefa.")

    update_data = ticket_data.dict(exclude_unset=True)
    if "description" in update_data and (not update_data["description"] or update_data["description"].strip() == ""):
        raise HTTPException(status_code=400, detail="A descrição da tarefa não pode estar vazia.")

    old_status = ticket.status
    old_priority = ticket.priority
    old_assigned_id = ticket.assigned_to_id
    old_description = ticket.description
    
    target_status = update_data.get("status")
    target_blocked = update_data.get("blocked_by_id", ticket.blocked_by_id)
    
    if target_blocked and target_status and target_status.lower() in ['in progress', 'em progresso', 'in review', 'em revisão']:
        prereq = db.query(Ticket).filter(Ticket.id == target_blocked).first()
        if prereq and prereq.status.lower() not in ['done', 'concluído', 'concluido']:
            raise HTTPException(
                status_code=400,
                detail=f"⚠️ Esta tarefa depende da conclusão da tarefa #{prereq.id} ('{prereq.title}') e ainda não pode ser iniciada."
            )

    for key, value in update_data.items():
        setattr(ticket, key, value)
        
    db.commit()
    db.refresh(ticket)
    
    changes = []
    if "status" in update_data and old_status != ticket.status:
        changes.append(f"Mudou o estado de '{old_status or 'N/D'}' para '{ticket.status}'")
        
    if "priority" in update_data and old_priority != ticket.priority:
        changes.append(f"Alterou a prioridade de '{old_priority}' para '{ticket.priority}'")

    if "description" in update_data and old_description != ticket.description:
        changes.append("Atualizou a descrição da tarefa")

    if "assigned_to_id" in update_data and old_assigned_id != ticket.assigned_to_id:
        new_assignee = db.query(User).filter(User.id == ticket.assigned_to_id).first() if ticket.assigned_to_id else None
        assignee_name = new_assignee.name if new_assignee and new_assignee.name else (new_assignee.email if new_assignee else "Ninguém")
        changes.append(f"Atribuiu a tarefa a: {assignee_name}")
        
        if ticket.assigned_to_id and ticket.assigned_to_id != current_user.id:
            db.add(Notification(
                user_id=ticket.assigned_to_id,
                message=f"A tarefa #{ticket.id} ('{ticket.title}') foi-te atribuída."
            ))
            db.commit()
        
    if not changes:
        changes.append("Atualizou os detalhes gerais da tarefa")
        
    detail_message = " | ".join(changes)
    log_action(db, current_user.id, "Atualização de Tarefa", detail_message, ticket_id=ticket.id)

    return ticket

@router.delete("/{ticket_id}", status_code=204)
def delete_ticket(
    ticket_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket).filter(Ticket.id == ticket_id)
    query = filter_tickets_by_permissions(query, current_user, db)
    db_ticket = query.first()
    
    if not db_ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada ou não tens permissão para apagá-la.")
    
    log_action(db, current_user.id, "Eliminação de Tarefa", f"Apagou a tarefa #{db_ticket.id} - {db_ticket.title}", ticket_id=db_ticket.id)
    db.delete(db_ticket)
    db.commit()
    return None

@router.post("/{ticket_id}/generate-ai-report")
def generate_ai_report(
    ticket_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket).filter(Ticket.id == ticket_id)
    query = filter_tickets_by_permissions(query, current_user, db)
    ticket = query.first()
    
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
    
    comments_text = "Nenhuma observação registada."
    try:
        comments = db.query(Comment).filter(Comment.ticket_id == ticket.id).all()
        if comments:
            comments_text = "\n".join([f"- {c.text}" for c in comments])
    except Exception:
        pass
    
    prompt = f"""
    És um assistente inteligente para técnicos de campo da empresa RFS. 
    Com base nos dados desta tarefa e nas observações de campo registadas, redige um resumo profissional e um relatório detalhado de intervenção em língua portuguesa.
    
    Título da Tarefa: {ticket.title}
    Descrição Inicial: {ticket.description or 'N/A'}
    Prioridade: {ticket.priority}
    
    Observações / Notas recolhidas durante a execução:
    {comments_text}
    """

    try:
        response = client.models.generate_content(
            model='gemini-flash-latest',
            contents=prompt,
        )
        return {"generated_report": response.text}
    except Exception as e:
        fallback_report = f"""RELATÓRIO DE INTERVENÇÃO TÉCNICA (RFS)

- Tarefa Executada: {ticket.title}
- Descrição Inicial: {ticket.description or 'Intervenção técnica regular efetuada conforme planeamento.'}
- Estado de Conclusão: Concluído com sucesso.
- Observações de Campo: {comments_text}"""
        
        return {"generated_report": fallback_report}

@router.get("/ai-focus-recommendation")
def get_ai_focus_recommendation(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket)
    query = filter_tickets_by_permissions(query, current_user, db)
    active_tickets = query.filter(Ticket.status != "Done").all()
    
    if not active_tickets:
        return {"recommendation": "De momento não tens tarefas ativas para analisar. Excelente trabalho!"}
        
    projects = db.query(Project).all()
    proj_map = {p.id: p.name for p in projects}

    today = date.today()
    today_str = today.isoformat()

    priority_weights = {'Crítica': 4, 'Alta': 3, 'Média': 2, 'Baixa': 1}

    scored_tasks = []
    for t in active_tickets:
        proj_name = proj_map.get(t.project_id, "Projeto Geral")
        due_info = str(t.due_date).split('T')[0] if t.due_date else None
        
        urgency_score = priority_weights.get(t.priority, 1) * 2  
        
        temporal_desc = "Sem prazo definido"
        if due_info:
            due_date_obj = date.fromisoformat(due_info)
            delta_days = (due_date_obj - today).days
            
            if delta_days < 0:
                urgency_score += abs(delta_days)
                temporal_desc = f"Atrasada por {abs(delta_days)} dia(s)"
            elif delta_days == 0:
                urgency_score += 15
                temporal_desc = "Vence EXATAMENTE HOJE"
            else:
                temporal_desc = "No prazo futuro"

        scored_tasks.append({
            "id": t.id,
            "title": t.title,
            "project": proj_name,
            "priority": t.priority,
            "due_date": due_info or "Sem prazo",
            "temporal": temporal_desc,
            "score": urgency_score
        })

    scored_tasks.sort(key=lambda x: x["score"], reverse=True)

    tasks_text = "\n".join([
        f"- [ID: #{st['id']}] Título: '{st['title']}' | Projeto: '{st['project']}' | Prioridade: {st['priority']} | Prazo: {st['due_date']} ({st['temporal']}) | Score Calculado: {st['score']}"
        for st in scored_tasks
    ])
    
    prompt = f"""
    A data atual é {today_str}. És o gestor de operações sénior da RFS.
    Abaixo tens a lista de tarefas ativas ordenada pelo sistema com base num algoritmo de cruzamento entre prioridade e prazos (Score de Urgência).

    Analisa criticamente os dados, valida qual é a tarefa mais crítica para o técnico executar agora (por exemplo, ponderando se uma tarefa que acaba hoje com prioridade Crítica supera uma tarefa atrasada de prioridade Baixa), e redige uma recomendação de foco altamente profissional, direta e justificada.

    Tarefas ativas:
    {tasks_text}
    """

    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt,
        )
        return {"recommendation": response.text} 
    except Exception as e:
        print(f"AVISO: IA de foco falhou. Erro: {str(e)}")
        top_task = scored_tasks[0]
        return {"recommendation": f"Análise Operacional Automática: Deves focar-te na tarefa #{top_task['id']} ('{top_task['title']}') devido ao cruzamento da sua prioridade ({top_task['priority']}) com o prazo ({top_task['due_date']})."}

@router.put("/{ticket_id}/complete", response_model=TicketResponse)
def complete_ticket(
    ticket_id: int,
    final_description: Optional[str] = Form(None),
    tracked_hours: Optional[float] = Form(0.0),
    file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket).filter(Ticket.id == ticket_id)
    query = filter_tickets_by_permissions(query, current_user, db)
    db_ticket = query.first()
    
    if not db_ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
    
    uncompleted_subtasks = [sub for sub in db_ticket.sub_tasks if not sub.is_completed]
    if uncompleted_subtasks:
        raise HTTPException(
            status_code=400, 
            detail=f"⚠️ Não podes concluir esta tarefa principal! Ainda existem {len(uncompleted_subtasks)} subtarefa(s) por concluir."
        )

    db_ticket.final_description = final_description or ""
    db_ticket.status = "Done"
    db_ticket.is_running = False
    
    if tracked_hours is not None and tracked_hours > 0:
        new_log = TimeLog(
            ticket_id=db_ticket.id,
            user_id=current_user.id,
            date=date.today(),
            hours_spent=float(tracked_hours)
        )
        db.add(new_log)

    if file and file.filename:
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        db_ticket.attachment_path = f"uploads/{file.filename}"

    db.commit()
    db.refresh(db_ticket)
    
    log_action(db, current_user.id, "Conclusão de Tarefa", f"Concluiu a tarefa #{db_ticket.id} - {db_ticket.title}", ticket_id=db_ticket.id)
    return db_ticket

def check_and_create_deadline_notifications(user: User, db: Session):
    try:
        today = date.today()
        today_str = today.isoformat()
        current_hour = datetime.now().hour
        
        query = db.query(Ticket)
        query = filter_tickets_by_permissions(query, user, db)
        tickets = query.filter(Ticket.status != "Done").all()
        
        for t in tickets:
            if not t.due_date:
                continue
            if isinstance(t.due_date, str):
                due_date_str = t.due_date.split('T')[0]
            elif hasattr(t.due_date, 'isoformat'):
                due_date_str = t.due_date.isoformat().split('T')[0]
            else:
                continue
            
            if due_date_str == today_str:
                msg = f"O prazo da tarefa '#{t.id} - {t.title}' acaba hoje!"
                exists = db.query(Notification).filter(Notification.user_id == user.id, Notification.message == msg).first()
                if not exists:
                    db.add(Notification(user_id=user.id, message=msg))
            elif due_date_str < today_str:
                msg = f"ATENÇÃO: A tarefa '#{t.id} - {t.title}' está atrasada!"
                exists = db.query(Notification).filter(Notification.user_id == user.id, Notification.message == msg).first()
                if not exists:
                    db.add(Notification(user_id=user.id, message=msg))
                    
        if current_hour >= 18:
            running_tickets = db.query(Ticket).filter(
                Ticket.assigned_to_id == user.id,
                Ticket.is_running == True
            ).all()
            
            for rt in running_tickets:
                cron_msg = f"ALERTA: Deixaste o cronómetro ligado na tarefa '#{rt.id} - {rt.title}' após o horário de expediente!"
                
                exists_cron = db.query(Notification).filter(
                    Notification.user_id == user.id, 
                    Notification.message == cron_msg,
                    Notification.created_at >= datetime.combine(today, datetime.min.time())
                ).first()
                
                if not exists_cron:
                    db.add(Notification(user_id=user.id, message=cron_msg))

        db.commit()
    except Exception as e:
        db.rollback()

@router.post("/my-day/generate-ai")
def generate_daily_ai_report(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    today = date.today()
    logs = db.query(TimeLog).filter(TimeLog.user_id == current_user.id, TimeLog.date == today).all()
    
    if not logs:
        return {
            "summary": "Sem tarefas registadas hoje.", 
            "detailed_report": "Não foram contabilizadas horas em nenhuma tarefa no dia de hoje."
        }
        
    hours_per_ticket = {}
    for log in logs:
        hours_per_ticket[log.ticket_id] = hours_per_ticket.get(log.ticket_id, 0) + log.hours_spent
        
    ticket_ids = list(hours_per_ticket.keys())
    tickets = db.query(Ticket).filter(Ticket.id.in_(ticket_ids)).all()

    tasks_info = []
    for t in tickets:
        dur_formatada = format_to_hhmm(hours_per_ticket[t.id])
        tasks_info.append(f"- Tarefa #{t.id}: {t.title} (Estado: {t.status}, Descrição: {t.description or 'N/A'}) - Tempo hoje: {dur_formatada}")
    
    tasks_text = "\n".join(tasks_info)
    
    prompt = f"""
    És um assistente técnico inteligente da RFS. Com base nas seguintes tarefas executadas pelo técnico no dia de hoje, gera:
    1. Um "summary" (um resumo curto e profissional de uma linha do dia).
    2. Um "detailed_report" (um relatório detalhado estruturado com as intervenções efetuadas, mantendo exatamente a formatação de tempo indicadas nas tarefas).
    
    Tarefas de hoje:
    {tasks_text}
    """

    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt,
        )
        text_result = response.text
        
        return {
            "summary": text_result[:150] + "...", 
            "detailed_report": text_result
        }
    except Exception as e:
        print(f"AVISO: IA falhou, a usar relatório padrão. Erro: {str(e)}")
        
        fallback_summary = f"Executadas {len(tickets)} intervenções técnicas planeadas para o dia de hoje."
        fallback_details = "RELATÓRIO DE ATIVIDADE DIÁRIA (RFS)\n\nIntervenções efetuadas:\n" + "\n".join(tasks_info)
        
        return {
            "summary": fallback_summary,
            "detailed_report": fallback_details
        }

@router.put("/my-day/today")
def update_daily_report(
    target_date: Optional[str] = Query(None),
    summary: Optional[str] = Form(None),
    detailed_report: Optional[str] = Form(None),
    kilometers: Optional[float] = Form(0.0),
    overtime_hours: Optional[float] = Form(0.0),
    pending_work: Optional[str] = Form(None),
    incidents: Optional[str] = Form(None),
    materials: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None), 
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    if target_date:
        today = datetime.strptime(target_date, "%Y-%m-%d").date()
    else:
        today = date.today()
        
    report = db.query(DailyReport).filter(DailyReport.user_id == current_user.id, DailyReport.date == today).first()
    
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado.")
        
    report.summary = summary
    report.detailed_report = detailed_report
    report.kilometers = kilometers
    report.overtime_hours = overtime_hours
    if hasattr(report, 'pending_work'): report.pending_work = pending_work
    if hasattr(report, 'incidents'): report.incidents = incidents
    if hasattr(report, 'materials'): report.materials = materials

    if file and file.filename:
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        report.image_path = f"uploads/{file.filename}"

    report.status = "Submetido" 
    report.rejection_reason = None
    report.submitted_at = datetime.now()
    
    db.commit()
    db.refresh(report)
    return {"message": "Relatório submetido com sucesso!"}

@router.put("/my-day/reopen")
def reopen_daily_report(
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    today = date.today()
    report = db.query(DailyReport).filter(DailyReport.user_id == current_user.id, DailyReport.date == today).first()
    
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado.")
        
    report.status = "Rascunho"
    db.commit()
    db.refresh(report)
    return {"message": "Relatório reaberto com sucesso!"}

@router.put("/admin/reports/{report_id}/status")
def update_report_status_admin(
    report_id: int,
    status_data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    role = getattr(current_user, "role", "Member")
    
    if role not in ["Admin", "Manager"]:
        raise HTTPException(status_code=403, detail="Acesso restrito. Apenas Admins e Managers podem aprovar relatórios.")
        
    report = db.query(DailyReport).filter(DailyReport.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado.")
        
    new_status = status_data.get("status")
    reason = status_data.get("rejection_reason") 

    if new_status == "Recusado":
        report.status = "Rascunho"
        report.rejection_reason = reason 
        
        data_formatada = report.date.strftime("%d/%m/%Y") if hasattr(report.date, 'strftime') else str(report.date)
        
        notif_msg = f"O teu relatório do dia {data_formatada} foi recusado. Motivo: {reason}"
        nova_notificacao = Notification(user_id=report.user_id, message=notif_msg)
        db.add(nova_notificacao)
        
    else:
        report.status = new_status
        if new_status == "Validado":
            report.rejection_reason = None 
            
    db.commit()
    db.refresh(report)
    return {"message": "Estado atualizado com sucesso!"}

@router.get("/{ticket_id}/comments")
def get_ticket_comments(
    ticket_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket).filter(Ticket.id == ticket_id)
    query = filter_tickets_by_permissions(query, current_user, db)
    ticket = query.first()
    
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
        
    comments = db.query(Comment).filter(Comment.ticket_id == ticket_id).all()
    
    enriched_comments = []
    for c in comments:
        author = db.query(User).filter(User.id == c.author_id).first()
        author_name = author.name if author and author.name else (author.email if author else "Colaborador")
        
        enriched_comments.append({
            "id": c.id,
            "text": c.text,
            "author_id": c.author_id,
            "author_name": author_name,
            "created_at": getattr(c, "created_at", None)
        })
        
    return enriched_comments

@router.post("/{ticket_id}/comments")
def create_ticket_comment(
    ticket_id: int,
    comment_data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(Ticket).filter(Ticket.id == ticket_id)
    query = filter_tickets_by_permissions(query, current_user, db)
    ticket = query.first()
    
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
        
    new_comment = Comment(
        text=comment_data.get("text"),
        ticket_id=ticket_id,
        author_id=current_user.id
    )
    db.add(new_comment)
    db.commit()
    db.refresh(new_comment)
    return new_comment

@router.get("/my-day/export-pdf")
def export_daily_report_pdf(
    target_date: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if target_date:
        report_date = datetime.strptime(target_date, "%Y-%m-%d").date()
    else:
        report_date = date.today()
        
    report = db.query(DailyReport).filter(
        DailyReport.user_id == current_user.id,
        DailyReport.date == report_date
    ).first()
    
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado para esta data.")
        
    logs = db.query(TimeLog).filter(
        TimeLog.user_id == current_user.id,
        TimeLog.date == report_date
    ).order_by(TimeLog.start_time.asc()).all()

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 40, "Relatório Diário de Atividade")
    
    p.setFont("Helvetica", 10)
    p.setFillColorRGB(0.3, 0.3, 0.3)
    p.drawString(50, height - 60, f"Técnico: {current_user.name or current_user.email}")
    p.drawString(50, height - 74, f"Data: {report.date} | Estado: {report.status}")
    p.setFillColorRGB(0, 0, 0)

    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, height - 105, "Registos de Trabalho (Horário da Escola - Por Sessão)")
    
    y_pos = height - 125
    p.setFont("Helvetica-Bold", 9)
    p.setFillColorRGB(0.2, 0.2, 0.2)
    p.drawString(50, y_pos, "ID")
    p.drawString(85, y_pos, "TAREFA")
    p.drawString(280, y_pos, "INÍCIO")
    p.drawString(350, y_pos, "FIM")
    p.drawString(430, y_pos, "DURAÇÃO")
    
    y_pos -= 5
    p.setStrokeColorRGB(0.7, 0.7, 0.7)
    p.line(50, y_pos, width - 50, y_pos)
    
    y_pos -= 15
    p.setFont("Helvetica", 9)
    p.setFillColorRGB(0, 0, 0)
    
    total_duration = 0.0
    for log in logs:
        t = db.query(Ticket).filter(Ticket.id == log.ticket_id).first()
        if not t:
            continue
            
        start_str = log.start_time.strftime("%H:%M") if getattr(log, "start_time", None) else "--:--"
        end_str = log.end_time.strftime("%H:%M") if getattr(log, "end_time", None) else "--:--"
        dur_float = round(log.hours_spent, 2)
        total_duration += dur_float
        
        p.drawString(50, y_pos, str(t.id))
        p.drawString(85, y_pos, str(t.title[:35]))
        p.drawString(280, y_pos, start_str)
        p.drawString(350, y_pos, end_str)
        p.drawString(430, y_pos, format_to_hhmm(dur_float))
        
        y_pos -= 15
        if y_pos < 100:
            break

    y_pos -= 5
    p.line(50, y_pos, width - 50, y_pos)
    y_pos -= 15
    p.setFont("Helvetica-Bold", 10)
    p.drawString(50, y_pos, f"Total Concluído: {format_to_hhmm(total_duration)} | Quilómetros: {report.kilometers or 0} km | Horas Extra: {report.overtime_hours or 0}h")

    y_pos -= 35
    p.setFont("Helvetica-Bold", 11)
    p.drawString(50, y_pos, "Resumo do Dia")
    
    y_pos -= 15
    p.setFont("Helvetica", 9)
    summary_text = report.summary or "Sem resumo registado."
    for line in summary_text.split('\n'):
        p.drawString(50, y_pos, line[:100])
        y_pos -= 12
        
    y_pos -= 10
    p.setFont("Helvetica-Bold", 11)
    p.drawString(50, y_pos, "Relatório Detalhado")
    
    y_pos -= 15
    p.setFont("Helvetica", 9)
    detailed_text = report.detailed_report or "Sem relatório detalhado."
    for line in detailed_text.split('\n'):
        clean_line = line.replace('■', '-').replace('\r', '')
        p.drawString(50, y_pos, clean_line[:100])
        y_pos -= 12
        if y_pos < 40:
            break

    p.showPage()
    p.save()
    
    buffer.seek(0)
    filename = f"Relatorio_{report.date}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/my-day/export-word")
def export_daily_report_word(
    target_date: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if target_date:
        report_date = datetime.strptime(target_date, "%Y-%m-%d").date()
    else:
        report_date = date.today()
        
    report = db.query(DailyReport).filter(
        DailyReport.user_id == current_user.id,
        DailyReport.date == report_date
    ).first()
    
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado para esta data.")
        
    doc = Document()
    doc.add_heading("Relatório Diário de Atividade", level=1)
    doc.add_paragraph(f"Técnico: {current_user.name or current_user.email}")
    doc.add_paragraph(f"Data: {report.date}")
    doc.add_paragraph(f"Resumo: {report.summary or 'Sem resumo.'}")
    doc.add_paragraph(f"Relatório Detalhado: {report.detailed_report or 'Sem detalhes.'}")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Relatorio_{report.date}.docx"}
    )

@router.put("/{ticket_id}/grab", response_model=TicketResponse)
def grab_ticket(
    ticket_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")
        
    if ticket.blocked_by_id:
        prereq = db.query(Ticket).filter(Ticket.id == ticket.blocked_by_id).first()
        if prereq and prereq.status.lower() not in ['done', 'concluído', 'concluido']:
            raise HTTPException(
                status_code=400,
                detail=f"Não podes agarrar esta tarefa! A dependência #{prereq.id} ('{prereq.title}') ainda não está concluída."
            )

    if ticket.assigned_to_id is not None and ticket.assigned_to_id != current_user.id:
        raise HTTPException(
            status_code=400, 
            detail="Tarde demais! Esta tarefa já foi atribuída a outro colega."
        )
        
    ticket.assigned_to_id = current_user.id
    ticket.status = "In Progress"
    
    db.commit()
    db.refresh(ticket)
    log_action(db, current_user.id, "Agarrar Tarefa", f"Agarrou a tarefa #{ticket.id}", ticket_id=ticket.id)
    return ticket

@router.post("/{ticket_id}/stop-timer", response_model=TicketResponse)
def stop_timer(
    ticket_id: int,
    timer_data: dict, 
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")

    ticket.is_running = False
    if "tracked_hours" in timer_data:
        ticket.tracked_hours = timer_data["tracked_hours"]

    start_str = timer_data.get("start_time")
    end_str = timer_data.get("end_time")
    
    start_dt = datetime.fromisoformat(start_str.replace("Z", "+00:00")) if start_str else None
    end_dt = datetime.fromisoformat(end_str.replace("Z", "+00:00")) if end_str else None

    new_log = TimeLog(
        ticket_id=ticket.id,
        user_id=current_user.id,
        date=date.today(),
        hours_spent=timer_data.get("session_hours", 0.0),
        start_time=start_dt,
        end_time=end_dt
    )
    
    db.add(new_log)
    db.commit()
    db.refresh(ticket)
    return ticket

# ==========================================
# GESTÃO DE SUBTAREFAS (COM ATRIBUIÇÃO)
# ==========================================

@router.get("/{ticket_id}/subtasks")
def get_subtasks(ticket_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    subtasks = db.query(SubTask).filter(SubTask.ticket_id == ticket_id).all()
    return [{"id": s.id, "title": s.title, "is_completed": s.is_completed, "assigned_to_id": getattr(s, "assigned_to_id", None)} for s in subtasks]

@router.post("/{ticket_id}/subtasks")
def create_subtask(ticket_id: int, data: dict, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    title = data.get("title")
    assigned_to_id = data.get("assigned_to_id")
    if not title or title.strip() == "":
        raise HTTPException(status_code=400, detail="O título da subtarefa é obrigatório.")
    
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa principal não encontrada.")
        
    sub = SubTask(ticket_id=ticket_id, title=title, is_completed=False, assigned_to_id=assigned_to_id)
    db.add(sub)
    
    if assigned_to_id and assigned_to_id != current_user.id:
        notif = Notification(
            user_id=assigned_to_id,
            message=f"Foi-te atribuída a subtarefa '{title}' na tarefa #{ticket.id} - {ticket.title}"
        )
        db.add(notif)

    db.commit()
    db.refresh(sub)
    return {"id": sub.id, "title": sub.title, "is_completed": sub.is_completed, "assigned_to_id": sub.assigned_to_id}

@router.put("/subtasks/{subtask_id}")
def update_subtask(subtask_id: int, data: dict, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    sub = db.query(SubTask).filter(SubTask.id == subtask_id).first()
    if not sub:
        raise HTTPException(status_code=404, detail="Subtarefa não encontrada.")
        
    old_assigned_id = getattr(sub, "assigned_to_id", None)

    if "is_completed" in data:
        sub.is_completed = bool(data["is_completed"])
    if "title" in data:
        sub.title = data["title"]
    if "assigned_to_id" in data:
        sub.assigned_to_id = data["assigned_to_id"]
        if sub.assigned_to_id and sub.assigned_to_id != old_assigned_id and sub.assigned_to_id != current_user.id:
            ticket = db.query(Ticket).filter(Ticket.id == sub.ticket_id).first()
            task_title = ticket.title if ticket else f"#{sub.ticket_id}"
            db.add(Notification(
                user_id=sub.assigned_to_id,
                message=f"Foi-te atribuída a subtarefa '{sub.title}' na tarefa {task_title}"
            ))
        
    db.commit()
    db.refresh(sub)
    return {"id": sub.id, "title": sub.title, "is_completed": sub.is_completed, "assigned_to_id": sub.assigned_to_id}

@router.delete("/subtasks/{subtask_id}", status_code=204)
def delete_subtask(subtask_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    sub = db.query(SubTask).filter(SubTask.id == subtask_id).first()
    if not sub:
        raise HTTPException(status_code=404, detail="Subtarefa não encontrada.")
        
    db.delete(sub)
    db.commit()
    return None

# ==========================================
# TIPOS DE TAREFA E ADMIN RELATÓRIOS
# ==========================================

from pydantic import BaseModel

class TaskTypeCreateSchema(BaseModel):
    name: str

class TaskTypeUpdateSchema(BaseModel):
    new_name: str

@router.get("/task-types/list")
def get_global_task_types(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_types = db.query(TaskType).all()
    return [{"id": t.id, "name": t.name} for t in db_types]

@router.post("/task-types/create")
def create_global_task_type(
    task_type: TaskTypeCreateSchema,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    role = getattr(current_user, "role", "Member")
    if role not in ["Admin", "Manager"]:
        raise HTTPException(status_code=403, detail="Apenas Managers ou Admins podem criar tipos de tarefa.")
    
    normalized_name = task_type.name.strip().capitalize()
    
    existing = db.query(TaskType).filter(TaskType.name == normalized_name).first()
    if existing:
        return {"message": "Tipo já existe", "name": normalized_name}
    
    new_type = TaskType(name=normalized_name)
    db.add(new_type)
    db.commit()
    db.refresh(new_type)
    return {"message": "Tipo criado com sucesso", "name": new_type.name}

@router.delete("/task-types/{type_id}", status_code=204)
def delete_global_task_type(
    type_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    role = getattr(current_user, "role", "Member")
    if role not in ["Admin", "Manager"]:
        raise HTTPException(status_code=403, detail="Apenas Managers ou Admins podem apagar tipos de tarefa.")
        
    db_type = db.query(TaskType).filter(TaskType.id == type_id).first()
    if not db_type:
        raise HTTPException(status_code=404, detail="Tipo de tarefa não encontrado.")
        
    db.delete(db_type)
    db.commit()
    return None

@router.get("/admin/reports/users-status")
def get_admin_users_reports_status(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    role = getattr(current_user, "role", "Member")
    if role not in ["Admin", "Manager"]:
        raise HTTPException(status_code=403, detail="Acesso negado.")

    users = db.query(User).all()
    result = []
    
    for u in users:
        reports = db.query(DailyReport).filter(
            DailyReport.user_id == u.id,
            DailyReport.status.in_(["Submetido", "Validado"])
        ).all()
        
        if not reports:
            continue
        
        reports_list = []
        for r in reports:
            reports_list.append({
                "id": r.id,
                "date": r.date.isoformat() if hasattr(r.date, 'isoformat') else str(r.date),
                "status": r.status,
                "summary": r.summary or "",
                "detailed_report": r.detailed_report or "",
                "kilometers": r.kilometers or 0.0,
                "overtime_hours": r.overtime_hours or 0.0,
                "image_path": getattr(r, "image_path", None),
                "submitted_at": r.submitted_at.isoformat() if getattr(r, "submitted_at", None) else None,
                "rejection_reason": getattr(r, "rejection_reason", None)
            })

        result.append({
            "user_id": u.id,
            "name": u.name or u.email,
            "email": u.email,
            "role": getattr(u, "role", "Member"),
            "reports": reports_list
        })

    return result

@router.post("/{ticket_id}/return")
def return_ticket(ticket_id: int, data: dict, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    reason = data.get("reason")
    if not reason or reason.strip() == "":
        raise HTTPException(status_code=400, detail="É obrigatório indicar um motivo/descrição para devolver a tarefa.")
    
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
        
    ticket.status = "To Do"
    ticket.assigned_to_id = None
    ticket.is_running = False
    ticket.return_reason = reason
    
    db.commit()
    log_action(db, current_user.id, "Devolução de Tarefa (Divórcio)", f"Devolveu a tarefa #{ticket.id}. Motivo: {reason}", ticket_id=ticket.id)
    return {"success": True, "message": "Tarefa devolvida com sucesso ao estado inicial."}

@router.post("/{ticket_id}/grab-team-task", response_model=TicketResponse)
def grab_team_task(
    ticket_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
        
    if ticket.assigned_to_id is not None:
        raise HTTPException(status_code=400, detail="Esta tarefa já foi agarrada por outro elemento.")

    if ticket.blocked_by_id:
        prereq = db.query(Ticket).filter(Ticket.id == ticket.blocked_by_id).first()
        if prereq and prereq.status.lower() not in ['done', 'concluído', 'concluido']:
            raise HTTPException(status_code=400, detail="A tarefa antecedente ainda não está concluída.")

    ticket.assigned_to_id = current_user.id
    ticket.status = "In Progress"
    
    db.commit()
    db.refresh(ticket)
    log_action(db, current_user.id, "Agarrar Tarefa de Equipa", f"Puxou para si a tarefa #{ticket.id} - {ticket.title}", ticket_id=ticket.id)
    return ticket


# ==========================================
# ENDPOINT DE AUDIT LOGS DA TAREFA ESPECÍFICA
# ==========================================

@router.get("/{ticket_id}/audit-logs")
def get_ticket_audit_logs(
    ticket_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    ticket = db.query(Ticket).filter(Ticket.id == ticket_id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")
        
    logs = db.query(AuditLog).filter(AuditLog.ticket_id == ticket_id).order_by(AuditLog.created_at.desc()).all()
    
    lista_logs = []
    for log in logs:
        user_obj = db.query(User).filter(User.id == log.user_id).first()
        username = user_obj.name if user_obj and user_obj.name else (user_obj.email if user_obj else f"User #{log.user_id}")

        lista_logs.append({
            "id": log.id,
            "user_id": log.user_id,
            "username": username,
            "project_id": getattr(log, "project_id", None),
            "ticket_id": getattr(log, "ticket_id", None),
            "action": log.action,
            "details": log.details,
            "created_at": log.created_at.isoformat() if log.created_at else None
        })
        
    return lista_logs